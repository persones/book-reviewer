print('book reviewer')
from docx import Document
import os

root_folder = 'C:/Users/perso/Google Drive/Shahar_Project Reliability Engineering/'
paths = ['Ready for Editorial Approval', 'Ready for Author Review']

def find_term(doc, *args):
  for p in doc.paragraphs:
    for term in args:
      pos = p.text.lower().find(term)
      if pos > -1:
        print(f'[{term}] {p.text[:40]}...{p.text[pos:pos + 40]}')

def find_style(doc, *args):
  for p in doc.paragraphs:
    found = False
    for r in p.runs:
      for term in args:
        if r.text.lower().find(term) > -1:
          print(f'[{term}] ({r.style.name}) {p.text[:70]}')
          found = True
       
def show_toc(doc, *argv):
  for p in doc.paragraphs:
    if p.style.name.find('Heading 1') > -1:
      print(p.text)

def run(func, *argv):
  for p in paths:
    print(f'---searching in {p}')
    folder = os.path.join(root_folder, p)
    files = os.listdir(folder) 
    for f in files:
      (name, ext) = os.path.splitext(f)
      if ext == '.docx':
        print(f'---Analyzing [{os.path.join(p, f)}]')
        try:
          doc = Document(os.path.join(folder, f)) 
        except:
          print(f'error analyizng file {f}')
          continue 
        func(doc, *argv)

#run(find_term, 'the case for o')
#run(find_style, 'the case for o')
run(show_toc, 0)